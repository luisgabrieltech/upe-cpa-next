import pandas as pd
import time
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import os
import re
from collections import Counter

# Mapeamento de IDs de questões para labels descritivos
qid_to_label = {
  "Questao 11": "Na GRADUAÇÃO, em qual curso você está matriculado(a)?",
  "Questao 12": "Qual o ano de ingresso no curso na UPE?",
  "Questao 13": "Em geral, durante o ano letivo, quantos dias por semana você vai para a UPE?",
  "Questao 14": "E em média, quantas horas por dia você permanece na UPE?",
  "Questao 15": "Em geral, quanto tempo você gasta de deslocamento entre a sua residência e a UPE (sua unidade de ensino)?",
  "Questao 16": "Qual forma de locomoção você mais utiliza para ir às aulas?",
  "Questao 17": {
    "Opcao 1": "Cursou o Ensino Médio em escola pública?",
    "Opcao 2": "Ingressou na UPE por meio de cota?",
    "Opcao 3": "Possui ou possuiu algum vínculo empregatício ainda no curso?"
  },
  "Questao 18": "Qual o seu nível de conhecimento sobre a existência e funcionamento da COMISSÃO PRÓPRIA DE AVALIAÇÃO (CPA), responsável pela avaliação interna da UPE?",
  "Questao 19": "Você já participou de alguma avaliação institucional promovida pela CPA?",
  "Questao 20": "Você teve conhecimento dos resultados GERAIS das avaliações institucionais anteriores promovidas pela Comissão Própria de Avaliação (CPA)?",
  "Questao 21": "Você teve conhecimento dos resultados das avaliações institucionais ESPECÍFICAS DA SUA UNIDADE promovidas pela Comissão Própria de Avaliação (CPA) e pelas Comissões Setoriais de Avaliação (CSA)?",
  "Questao 22": "Você identificou melhorias na UPE decorrentes de avaliações anteriores? Se sim, como você avalia essas melhorias?",
  "Questao 23": "Como você avalia as ações de divulgação e informação promovidas na sua unidade visando a participação dos estudantes na Autoavaliação Institucional promovida pela CPA?",
  "Questao 24": {
    "Opcao 1": "Fortalecimento da atuação para o desenvolvimento, com foco nas Políticas de Inovação, Sustentabilidade e Inclusão Social",
    "Opcao 2": "Promoção da diminuição das taxas de evasão e retenção",
    "Opcao 3": "Promoção da melhoria da qualidade dos cursos de graduação e pós-graduação",
    "Opcao 4": "Promoção da integração e da comunicação entre suas unidades",
    "Opcao 5": "Aumento, qualificação e promoção da Valorização do quadro docente, técnico e administrativo",
    "Opcao 6": "Otimização da gestão financeira, orçamentária e a captação de recursos"
  },
  "Questao 25": "Você tem conhecimento da missão da UPE?",
  "Questao 26": "Você, enquanto discente da UPE, considera que contribui para o cumprimento da missão institucional?",
  "Questao 27": "Enquanto discente, você considera que a missão institucional da UPE se reflete diretamente na formação acadêmica que a UPE lhe proporciona?",
  "Questao 28": {
    "Opcao 1": "Plano de Desenvolvimento Institucional (PDI)",
    "Opcao 2": "Estatuto",
    "Opcao 3": "Regimento Geral",
    "Opcao 4": "Guia (Manual) do Estudante UPE"
  },
  "Questao 29": "Como você avalia as ações efetivadas entre 2019 e 2024 do Plano de Desenvolvimento Institucional (PDI) da UPE?",
  "Questao 30": "Você contribuiu ou está contribuindo para a construção, discussão e avaliação do Novo Plano de Desenvolvimento Institucional (PDI) da UPE (2025/2029) que está em construção atualmente?",
  "Questao 31": "Na sua perspectiva, existe uma formulação clara dos objetivos e finalidades da UPE?",
  "Questao 32": "Você conhece a resolução CONSUN 017/2021 sobre a Política de Acessibilidade inclusão Educacional da UPE?",
  "Questao 33": "Como você vê a gestão da sua Unidade quanto às demandas de acessibilidade e inclusão da comunidade acadêmica?",
  "Questao 34": "Você já viu na UPE alguma orientação, campanha ou ação voltada ao tema acessibilidade?",
  "Questao 35": {
    "descricao": "Em termos de impacto social, como você avalia a viabilização das seguintes ações/órgãos da UPE?",
    "opcoes": {
      "Opcao 1": "Atendimento do Complexo Hospitalar (CISAM, PROCAPE, HUOC)",
      "Opcao 2": "Clínicas de atendimento odontológico (FOP)",
      "Opcao 3": "Escolas de Aplicação",
      "Opcao 4": "Instituto Confúcio",
      "Opcao 5": "Núcleo de Estudos sobre Violência e Promoção da Saúde da UPE",
      "Opcao 6": "Pré-vestibular da Universidade de Pernambuco (PREVUPE)",
      "Opcao 7": "Programa de ensino e pesquisa em emergências, acidentes e violências (PEPEAV)",
      "Opcao 8": "Programa de Línguas e Informática da UPE (PROLINFO)",
      "Opcao 9": "Núcleo de Diversidade e Identidades Sociais (NDIS)"
    }
  },
  "Questao 36": {
    "descricao": "UPE assume a Responsabilidade Social como um princípio. Como você avalia as ações realizadas para viabilização dos compromissos sociais contemplados no PDI?",
    "opcoes": {
      "Opcao 1": "A erradicação de todas as formas de exclusão social.",
      "Opcao 2": "A garantia do sistema de direitos coletivos e individuais.",
      "Opcao 3": "Respeito ao processo democrático no País, no Estado e na própria UPE.",
      "Opcao 4": "A universalização e elevação da qualidade da educação pública.",
      "Opcao 5": "O desenvolvimento sustentável articulando o crescimento humano com a preservação da natureza.",
      "Opcao 6": "As políticas de promoção da paz.",
      "Opcao 7": "A promoção da igualdade entre os sexos e a autonomia das mulheres.",
      "Opcao 8": "Garantir a sustentabilidade com qualidade de vida.",
      "Opcao 9": "Estabelecer uma Parceria Mundial para o Desenvolvimento."
    }
  },
  "Questao 37": {
    "descricao": "Considerando a responsabilidade social da UPE com a sociedade pernambucana, você identificou ações direcionadas à comunidade do entorno de sua Unidade, considerando os aspectos abaixo?",
    "opcoes": {
      "Opcao 1": "Ações sociais relacionadas à saúde, esporte, lazer e educação.",
      "Opcao 2": "Ações relacionadas à preservação e conservação ambiental.",
      "Opcao 3": "Ações relacionadas à produção artística e cultural."
    }
  }
}

GEMINI_API_KEY = "AIzaSyDolQ8860QT5M3yZnQ4D39RNFDRVMWWaNE"
RATE_LIMIT = 60 
MAX_REQUESTS = 15

def gerar_analise_gemini(questao, dados, request_count, total_respostas=None):
    if request_count > 0 and request_count % MAX_REQUESTS == 0:
        print(f"[DEBUG] Atingiu o limite de {MAX_REQUESTS} requisições. Aguardando {RATE_LIMIT} segundos...")
        time.sleep(RATE_LIMIT) 
        print("Retomando as requisições...")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={GEMINI_API_KEY}"

    # Obter o contexto da questão, se disponível
    contexto_questao = ""
    label_info = qid_to_label.get(questao, "")
    
    # Verificar se a questão tem subopções
    if isinstance(label_info, dict):
        if "descricao" in label_info:
            contexto_questao = label_info["descricao"]
            # Adicionar informações sobre as opções, se disponíveis
            if "opcoes" in label_info:
                contexto_questao += "\nOpções: "
                for opcao, descricao in label_info["opcoes"].items():
                    contexto_questao += f"\n- {opcao}: {descricao}"
        elif len(label_info) > 0:  # Para questões como a 17, 24, etc.
            contexto_questao = "Subquestões: "
            for opcao, descricao in label_info.items():
                contexto_questao += f"\n- {opcao}: {descricao}"
    else:
        contexto_questao = label_info
    
    prompt = f"""
    Analise os dados da seguinte pesquisa:
    Você é um analista de pesquisa da Universidade de Pernambuco e trabalha na Comissão Própria de Avaliação da faculdade.
    Você está analisando um questionário voltado para os alunos das unidades de ensino da Universidade de Pernambuco.

    **Pergunta:** {questao}
    **Contexto da pergunta:** {contexto_questao}
    **Total de respostas recebidas:** {total_respostas if total_respostas is not None else sum(dados.values())}

    **Distribuição das respostas:** {dados}

    Forneça uma interpretação objetiva e concisa, destacando:
    - A resposta mais frequente e seu percentual.
    - Possíveis implicações desses resultados.
    - Qualquer tendência ou ponto de atenção relevante.

    A análise deve ser direta, clara e sem floreios.
    Caso numero de respostas seja menor que 10, tente analisar a questão com outro ponto de vista mas alerte que o numero de respostas é pequeno.  
    """

    payload = {
        "contents": [
            {
                "parts": [{"text": prompt}]
            }
        ]
    }

    headers = {"Content-Type": "application/json"}

    response = requests.post(url, headers=headers, json=payload)

    if response.status_code == 429:
        print("[DEBUG] Limite de requisições atingido! Aguardando antes de tentar novamente...")
        time.sleep(RATE_LIMIT)
        return gerar_analise_gemini(questao, dados, request_count, total_respostas)

    response.raise_for_status()
    resposta_json = response.json()

    candidates = resposta_json.get("candidates", [])
    if candidates:
        content = candidates[0].get("content", {})
        parts = content.get("parts", [])
        if parts:
            return parts[0].get("text", "Erro: Sem texto gerado pela IA.")
    return "Erro: Sem resposta válida da IA."

def criar_grafico(questao, dados):
    plt.figure(figsize=(8, 5))

    bars = plt.barh(dados['Resposta'], dados['Frequência'], color='#1f77b4', edgecolor='black')

    max_value = dados['Frequência'].max()
    margem_extra = max_value * 0.15 

    total_respostas = dados['Frequência'].sum()
    for bar, valor in zip(bars, dados['Frequência']):
        percentual = (valor / total_respostas) * 100
        plt.text(bar.get_width() + (max_value * 0.02), 
                 bar.get_y() + bar.get_height()/2, 
                 f"{percentual:.1f}%", 
                 va='center', fontsize=10)

    plt.xlabel("")
    plt.ylabel("")
    plt.title("")
    plt.gca().invert_yaxis()
    plt.xticks([])  

    plt.xlim(0, max_value + margem_extra)

    for spine in plt.gca().spines.values():
        spine.set_visible(True)  
        spine.set_color("black")  
        spine.set_linewidth(1.5)  

    nome_arquivo = re.sub(r'[^a-zA-Z0-9]', '_', questao)[:50]

    arquivo_temp = f"grafico_{nome_arquivo}.png"
    try:
        plt.savefig(arquivo_temp, dpi=300, bbox_inches='tight')
    except Exception as e:
        print(f"[DEBUG] Erro ao salvar gráfico: {e}")
        arquivo_temp = f"grafico_{hash(questao)}.png"
        plt.savefig(arquivo_temp, dpi=300, bbox_inches='tight')
    
    plt.close()
    return arquivo_temp

def gerar_relatorio(df, saida_word):
    documento = Document()
    documento.add_heading("Relatório Analítico - Análise de Respostas", level=1)

    colunas = [col for col in df.columns if col.startswith('Questao')]
    colunas.sort()

    request_count = 0
    for coluna in colunas:
        if coluna not in df.columns or df[coluna].dropna().empty:
            print(f"Pulando {coluna}: Coluna não encontrada ou vazia")
            continue

        respostas = df[coluna].dropna().astype(str)
        total_respostas = len(respostas)
        contagem_respostas = Counter(respostas)
        dados_questao = pd.DataFrame(contagem_respostas.items(), columns=['Resposta', 'Frequência'])
        dados_questao = dados_questao.sort_values(by='Frequência', ascending=False)

        documento.add_heading(coluna, level=2)
        documento.add_paragraph(f"A questão '{coluna}' recebeu {total_respostas} respostas únicas.")

        if len(dados_questao) > 10:
            tabela = documento.add_table(rows=1, cols=3)
            tabela.style = 'Table Grid'
            hdr_cells = tabela.rows[0].cells
            hdr_cells[0].text = 'Resposta'
            hdr_cells[1].text = 'Frequência'
            hdr_cells[2].text = 'Percentual (%)'

            total_respostas_unicas = dados_questao['Frequência'].sum()
            for _, row in dados_questao.iterrows():
                percentual = (row['Frequência'] / total_respostas_unicas) * 100
                row_cells = tabela.add_row().cells
                row_cells[0].text = row['Resposta']
                row_cells[1].text = str(row['Frequência'])
                row_cells[2].text = f"{percentual:.1f}"
        else:
            grafico_path = criar_grafico(coluna, dados_questao)
            documento.add_picture(grafico_path, width=Inches(6.0))
            os.remove(grafico_path)

        analise_gemini = gerar_analise_gemini(coluna, contagem_respostas, request_count, total_respostas)
        request_count += 1
        documento.add_heading("Análise da IA", level=3)
        documento.add_paragraph(analise_gemini)

        documento.add_paragraph("Fonte: Autor (2024)", style='Caption')

    documento.save(saida_word)
    print(f"Relatório gerado: {saida_word}")

def carregar_dados(arquivo, aba=None):
    try:
        if aba is None:
            df = pd.read_excel(arquivo)
        else:
            df = pd.read_excel(arquivo, sheet_name=aba)
        print(f"Dados carregados com sucesso. Total de linhas: {len(df)}")
        return df
    except Exception as e:
        print(f"Erro ao carregar dados: {e}")
        raise

file_path = "respostas.xlsx"

df_respostas = carregar_dados(file_path)

saida_word = "relatorio_analiseIA.docx"
gerar_relatorio(df_respostas, saida_word)